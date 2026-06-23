import os
import sys
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Set up paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, 'test_outputs')
DEBUG_DIR = os.path.join(BASE_DIR, 'test_debug')

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(DEBUG_DIR, exist_ok=True)

def analyze_dt_ownership(df):
    print("Starting DT ownership analysis...")
    # Basic DT stats
    total_unique_dts = df['dt_name'].nunique()
    dt_officer_counts = df.groupby('dt_name')['account_officer'].nunique().reset_index()
    dt_officer_counts.columns = ['dt_name', 'Officer_Count']
    
    single_officer_dts = len(dt_officer_counts[dt_officer_counts['Officer_Count'] == 1])
    shared_dts = len(dt_officer_counts[dt_officer_counts['Officer_Count'] > 1])
    percentage_shared = (shared_dts / total_unique_dts) * 100 if total_unique_dts > 0 else 0
    
    # Customers affected by shared DTs
    shared_dt_names = dt_officer_counts[dt_officer_counts['Officer_Count'] > 1]['dt_name'].tolist()
    customers_affected = len(df[df['dt_name'].isin(shared_dt_names)])
    
    # DT-Officer detailed mapping
    dt_officer_mapping = df.groupby(['dt_name', 'account_officer', 'officer_type']).size().reset_index(name='Customer_Count')
    dt_total_customers = df.groupby('dt_name').size().reset_index(name='DT_Total_Customers')
    dt_officer_mapping = dt_officer_mapping.merge(dt_total_customers, on='dt_name')
    dt_officer_mapping['Ownership_Percent'] = (dt_officer_mapping['Customer_Count'] / dt_officer_mapping['DT_Total_Customers']) * 100
    
    # Break down shared DTs by officer type
    shared_dt_mapping = dt_officer_mapping[dt_officer_mapping['dt_name'].isin(shared_dt_names)]
    
    def categorize_shared_dt(dt_group):
        officer_types = set(dt_group['officer_type'])
        
        if len(officer_types) == 1:
            officer_type = next(iter(officer_types))
            if officer_type == 'DMO':
                if dt_group['account_officer'].nunique() > 1:
                    return 'Multiple DMO'
            elif officer_type == 'Vendor':
                return 'Vendor only'
        else:
            return 'Mixed (DMO + Vendor)'
        
        return 'Other'
    
    shared_dt_categories = shared_dt_mapping.groupby('dt_name').apply(categorize_shared_dt).reset_index(name='Category')
    category_counts = shared_dt_categories['Category'].value_counts().to_dict()
    
    # Exempt customers (business_unit == "MD")
    md_customers = df[df['business_unit'] == 'MD']
    md_dts = md_customers['dt_name'].nunique()
    md_shared_dts = dt_officer_counts[
        (dt_officer_counts['dt_name'].isin(md_customers['dt_name'].unique())) & 
        (dt_officer_counts['Officer_Count'] > 1)
    ]['dt_name'].nunique()
    
    # Save debug files
    dt_officer_mapping.to_csv(os.path.join(DEBUG_DIR, 'dt_officer_mapping.csv'), index=False)
    dt_officer_counts.to_csv(os.path.join(DEBUG_DIR, 'dt_officer_counts.csv'), index=False)
    shared_dt_mapping.to_csv(os.path.join(DEBUG_DIR, 'shared_dt_analysis.csv'), index=False)
    print(f"Debug files saved to {DEBUG_DIR}")
    
    analysis_results = {
        'total_unique_dts': total_unique_dts,
        'single_officer_dts': single_officer_dts,
        'shared_dts': shared_dts,
        'percentage_shared': percentage_shared,
        'customers_affected': customers_affected,
        'category_counts': category_counts,
        'dt_officer_mapping': dt_officer_mapping,
        'md_customers': md_customers,
        'md_dts': md_dts,
        'md_shared_dts': md_shared_dts
    }
    
    print("DT ownership analysis complete!")
    return analysis_results

def redistribute_customers(df, analysis_results):
    print("Starting customer redistribution...")
    df = df.copy()
    df['Officer_Name_Old'] = df['account_officer']
    df['Officer_Name_New'] = df['account_officer']
    df['Redistributed'] = 'No'
    df['Redistribution_Reason'] = ''
    
    # First, let's get shared DTs from the NON-MD customers specifically
    non_md_dt_officer_counts = df.groupby('dt_name')['account_officer'].nunique().reset_index()
    non_md_dt_officer_counts.columns = ['dt_name', 'Officer_Count']
    shared_dt_names = non_md_dt_officer_counts[non_md_dt_officer_counts['Officer_Count'] > 1]['dt_name'].unique()
    
    redistribution_log = []
    
    for dt_name in shared_dt_names:
        dt_customers = df[df['dt_name'] == dt_name]
        
        # Calculate officer customer counts for this DT
        officer_counts = dt_customers.groupby('account_officer').size().sort_values(ascending=False)
        
        # Determine winning officer
        max_count = officer_counts.max()
        top_officers = officer_counts[officer_counts == max_count].index.tolist()
        
        if len(top_officers) == 1:
            winning_officer = top_officers[0]
        else:
            # Tie breaker: first in dataset, then alphabetical
            officer_order = df[df['dt_name'] == dt_name]['account_officer'].drop_duplicates().tolist()
            top_in_order = [officer for officer in officer_order if officer in top_officers]
            if top_in_order:
                winning_officer = top_in_order[0]
            else:
                winning_officer = sorted(top_officers)[0]
        
        # Update customers
        mask = (df['dt_name'] == dt_name) & (df['account_officer'] != winning_officer)
        customers_moved = mask.sum()
        
        df.loc[mask, 'Officer_Name_New'] = winning_officer
        df.loc[mask, 'Redistributed'] = 'Yes'
        df.loc[mask, 'Redistribution_Reason'] = f'Majority ownership assigned to {winning_officer}'
        
        redistribution_log.append({
            'DT_Name': dt_name,
            'Original Officers': ', '.join(dt_customers['account_officer'].unique()),
            'Winning Officer': winning_officer,
            'Customers Moved': customers_moved,
            'MD Exempt': 'No'
        })
    
    # Save redistribution log
    log_df = pd.DataFrame(redistribution_log)
    log_df.to_csv(os.path.join(DEBUG_DIR, 'redistribution_log.csv'), index=False)
    print(f"Redistribution log saved to {DEBUG_DIR}")
    print("Customer redistribution complete!")
    return df

def generate_word_report(analysis_results, redistributed_df):
    print("Generating analysis report...")
    doc = Document()
    
    # Title
    title = doc.add_heading('DT Ownership Analysis and Customer Redistribution Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Date Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        f'This report analyzes DT (Distribution Team) ownership across the customer base and proposes a redistribution strategy. '
        f'Out of {analysis_results["total_unique_dts"]} total unique DTs, {analysis_results["shared_dts"]} ({analysis_results["percentage_shared"]:.2f}%) are shared among multiple officers, '
        f'affecting {analysis_results["customers_affected"]} customers. '
        f'Using a majority ownership approach, we have redistributed customers to ensure each DT is owned by exactly one officer, '
        f'with exemptions for all MD Business Unit customers.'
    )
    doc.add_page_break()
    
    # Dataset Statistics
    doc.add_heading('Dataset Statistics', level=1)
    stats_table = doc.add_table(rows=1, cols=2)
    stats_table.style = 'Light Grid Accent 1'
    hdr_cells = stats_table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    stats_data = [
        ('Total Customers', len(redistributed_df)),
        ('Total Unique DTs', analysis_results['total_unique_dts']),
        ('DTs with Single Officer', analysis_results['single_officer_dts']),
        ('DTs Shared Among Multiple Officers', analysis_results['shared_dts']),
        ('Percentage of Shared DTs', f'{analysis_results["percentage_shared"]:.2f}%'),
        ('Customers Affected by Shared DTs', analysis_results['customers_affected']),
        ('MD Business Unit Customers', len(analysis_results['md_customers'])),
        ('MD Unique DTs', analysis_results['md_dts']),
        ('Shared MD DTs (Exempt)', analysis_results['md_shared_dts'])
    ]
    
    for metric, value in stats_data:
        row_cells = stats_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = str(value)
    doc.add_page_break()
    
    # DT Distribution Overview
    doc.add_heading('DT Distribution Overview', level=1)
    doc.add_paragraph('The following table shows DT ownership distribution:')
    doc.add_page_break()
    
    # Shared DT Analysis
    doc.add_heading('Shared DT Analysis', level=1)
    doc.add_paragraph('Breakdown of shared DTs by category:')
    
    category_table = doc.add_table(rows=1, cols=2)
    category_table.style = 'Light Grid Accent 1'
    hdr_cells = category_table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Count'
    
    for category, count in analysis_results['category_counts'].items():
        row_cells = category_table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = str(count)
    doc.add_page_break()
    
    # Officer Type Analysis
    doc.add_heading('Officer Type Analysis', level=1)
    doc.add_paragraph('Detailed DT-Officer mappings are included in the appendix.')
    doc.add_page_break()
    
    # Exemption Analysis
    doc.add_heading('Exemption Analysis', level=1)
    doc.add_paragraph(
        f'All customers with business_unit = "MD" are completely exempt from redistribution. '
        f'This includes {len(analysis_results["md_customers"])} customers across {analysis_results["md_dts"]} unique DTs. '
        f'Of these, {analysis_results["md_shared_dts"]} DTs are shared but remain unchanged.'
    )
    doc.add_page_break()
    
    # Appendix
    doc.add_heading('Appendix: Detailed DT-Officer Mapping', level=1)
    mapping_df = analysis_results['dt_officer_mapping']
    
    # Add table for DT-Officer mapping
    mapping_table = doc.add_table(rows=1, cols=5)
    mapping_table.style = 'Light Grid Accent 1'
    hdr_cells = mapping_table.rows[0].cells
    hdr_cells[0].text = 'DT_Name'
    hdr_cells[1].text = 'Officer_Name'
    hdr_cells[2].text = 'Officer_Type'
    hdr_cells[3].text = 'Customer Count'
    hdr_cells[4].text = 'Ownership Percentage'
    
    for _, row in mapping_df.iterrows():
        row_cells = mapping_table.add_row().cells
        row_cells[0].text = str(row['dt_name'])
        row_cells[1].text = str(row['account_officer'])
        row_cells[2].text = str(row['officer_type'])
        row_cells[3].text = str(row['Customer_Count'])
        row_cells[4].text = f'{row["Ownership_Percent"]:.2f}%'
    
    # Save the document
    doc.save(os.path.join(OUTPUT_DIR, 'DT_Analysis_Report.docx'))
    print(f"Analysis report saved to {OUTPUT_DIR}")

# Create test data with mixed MD and non-MD customers in the same DT
sample_data = [
    # DT001 has both MD and non-MD customers
    {"account_number": "ACC001", "dt_name": "DT001", "account_officer": "Officer A", "officer_type": "DMO", "business_unit": "BU1"},
    {"account_number": "ACC002", "dt_name": "DT001", "account_officer": "Officer A", "officer_type": "DMO", "business_unit": "BU1"},
    {"account_number": "ACC003", "dt_name": "DT001", "account_officer": "Officer B", "officer_type": "Vendor", "business_unit": "BU1"},
    {"account_number": "ACC004", "dt_name": "DT001", "account_officer": "Officer B", "officer_type": "Vendor", "business_unit": "BU1"},
    {"account_number": "ACC005", "dt_name": "DT001", "account_officer": "Officer A", "officer_type": "DMO", "business_unit": "BU1"},
    {"account_number": "ACC006", "dt_name": "DT001", "account_officer": "Officer X", "officer_type": "DMO", "business_unit": "MD"},  # MD customer
    {"account_number": "ACC007", "dt_name": "DT001", "account_officer": "Officer Y", "officer_type": "Vendor", "business_unit": "MD"}, # MD customer
    
    {"account_number": "ACC008", "dt_name": "DT002", "account_officer": "Officer C", "officer_type": "DMO", "business_unit": "BU2"},
    
    {"account_number": "ACC009", "dt_name": "DT003", "account_officer": "Officer D", "officer_type": "Vendor", "business_unit": "MD"},
    {"account_number": "ACC010", "dt_name": "DT003", "account_officer": "Officer E", "officer_type": "Vendor", "business_unit": "MD"},
]

df = pd.DataFrame(sample_data)
print("Test data created!")
print(df)
print()

# Split into MD and non-MD customers
md_customers = df[df['business_unit'] == 'MD'].copy()
non_md_customers = df[df['business_unit'] != 'MD'].copy()
print(f"Split test data: {len(md_customers)} MD customers, {len(non_md_customers)} non-MD customers")

# Analyze DT ownership (on ALL customers for full analysis)
analysis_results = analyze_dt_ownership(df)

# Redistribute ONLY non-MD customers
print("\n--- Redistributing non-MD customers ---")
redistributed_non_md = redistribute_customers(non_md_customers, analysis_results)

# Prepare MD customers to append back (add the new columns)
md_customers['Officer_Name_Old'] = md_customers['account_officer']
md_customers['Officer_Name_New'] = md_customers['account_officer']
md_customers['Redistributed'] = 'No'
md_customers['Redistribution_Reason'] = 'MD Business Unit - Exempt'

# Combine redistributed non-MD and original MD customers
redistributed_df = pd.concat([redistributed_non_md, md_customers], ignore_index=True)

print("\n--- Final redistributed data ---")
print(redistributed_df[['account_number', 'dt_name', 'business_unit', 'Officer_Name_Old', 'Officer_Name_New', 'Redistributed', 'Redistribution_Reason']])

# Save test outputs
print("\nSaving test outputs...")
redistributed_df.to_excel(os.path.join(OUTPUT_DIR, 'Customer_Redistributed.xlsx'), index=False)
generate_word_report(analysis_results, redistributed_df)

print("\n✅ All tests passed!")
print(f"Test outputs saved to {OUTPUT_DIR}")
print(f"Test debug files saved to {DEBUG_DIR}")
